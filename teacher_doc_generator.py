from docx import Document
from docx.shared import Inches
from odf import text, teletype
from odf.opendocument import load
import os
import json
from datetime import datetime
import pandas as pd
from PIL import Image
import shutil
from google.oauth2 import service_account
from googleapiclient.discovery import build
import io
from PyQt6.QtWidgets import (QApplication, QMainWindow, QWidget, QPushButton, 
                            QLabel, QVBoxLayout, QHBoxLayout, QLineEdit, 
                            QTabWidget, QScrollArea, QComboBox, QMessageBox,
                            QFileDialog, QFrame, QGridLayout)
from PyQt6.QtCore import Qt, QTimer
from PyQt6.QtGui import QFont
from googleapiclient.http import MediaIoBaseDownload

class LoginManager:
    def __init__(self, creds):
        self.SPREADSHEET_ID = '171kBtkTN5LUTNVMkCEhJ-L-l4xh7Wj9sJE76W0HmZ9w'
        self.creds = creds
        self.service = build('sheets', 'v4', credentials=creds)
    
    def verify_login(self, login_code):
        try:
            # 讀取登入碼（A欄
            result = self.service.spreadsheets().values().get(
                spreadsheetId=self.SPREADSHEET_ID,
                range='login!A:A'
            ).execute()
            login_codes = result.get('values', [])
            
            # 檢查登入碼是否存在
            if not any(code[0] == login_code for code in login_codes if code):
                return False, None
            
            # 找到對應的最後更新日期（C欄）
            for i, code in enumerate(login_codes):
                if code and code[0] == login_code:
                    row_num = i + 1
                    break
            
            # 讀取該使用者的最後更新日期
            result = self.service.spreadsheets().values().get(
                spreadsheetId=self.SPREADSHEET_ID,
                range=f'login!C{row_num}'
            ).execute()
            last_update = result.get('values', [[None]])[0][0]
            
            # 讀取師資和課程的最後更新日期
            result = self.service.spreadsheets().values().get(
                spreadsheetId=self.SPREADSHEET_ID,
                range='login!D2:E2'
            ).execute()
            data_update = result.get('values', [[None, None]])[0]
            
            need_update = False
            if last_update:
                try:
                    # 移除秒數部分，只保留到分鐘
                    last_update = last_update.split(':')[0] + ':' + last_update.split(':')[1]
                    last_update = datetime.strptime(last_update, '%Y-%m-%d %H:%M')
                    
                    if data_update[0]:  # 檢查師資更新日期
                        teacher_update = data_update[0].split(':')[0] + ':' + data_update[0].split(':')[1]
                        teacher_update = datetime.strptime(teacher_update, '%Y-%m-%d %H:%M')
                        if teacher_update > last_update:
                            need_update = True
                            
                    if data_update[1]:  # 檢查課程更新日期
                        course_update = data_update[1].split(':')[0] + ':' + data_update[1].split(':')[1]
                        course_update = datetime.strptime(course_update, '%Y-%m-%d %H:%M')
                        if course_update > last_update:
                            need_update = True
                except ValueError as e:
                    print(f"日期格式錯誤：{str(e)}")
                    need_update = True
            else:
                need_update = True
            
            # 更新最後更新日期（使用統一的格式）
            now = datetime.now().strftime('%Y-%m-%d %H:%M')
            self.service.spreadsheets().values().update(
                spreadsheetId=self.SPREADSHEET_ID,
                range=f'login!C{row_num}',
                valueInputOption='RAW',
                body={'values': [[now]]}
            ).execute()
            
            return True, need_update
            
        except Exception as e:
            print(f"驗證錯誤：{str(e)}")
            return False, None
class TeacherDocApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("教師資料文件產生器")
        self.setFixedSize(250, 150)
        
        # 設定視窗永遠在最上層
        self.setWindowFlags(self.windowFlags() | Qt.WindowType.WindowStaysOnTopHint)
        
        # 加入試算表 ID
        self.SPREADSHEET_ID = '171kBtkTN5LUTNVMkCEhJ-L-l4xh7Wj9sJE76W0HmZ9w'
        
        # 修改教師標記定義
        self.teacher_tags = [
            ("姓名", "@name"),
            ("綽號", "@nickname"),
            ("大頭照", "@photo"),
            ("申請單位", "@unit"),
            ("出生日期", "@birth"),
            ("性別", "@gender"),
            ("手機", "@mobile"),
            ("身分證字號", "@idno"),
            ("通訊地址", "@address"),
            ("Email", "@email"),
            ("Line ID", "@line"),
            ("專長", "@skill"),
            ("最高學歷", "@education"),
            ("現職", "@job"),
            ("教學經驗（年）", "@experience"),
            ("經歷", "@history"),
            ("身分證正面（照片）", "@id_front"),
            ("身分證反面（照片）", "@id_back"),
            ("畢業證書（照片）", "@diploma"),
            ("其他（良民證、比賽等）", "@other_certs")
        ]
        
        # 修改課程標記定義，分成三個區塊
        self.course_tags = [
            ("社團名稱", "@course_name"),
            ("課程介紹", "@intro"),
            ("教學目標", "@target"),
            ("材料費", "@material_fee"),
            ("材料內容", "@reason"),     
            ("課程主題（表格）", "@course_topic"),       
            ("課程內容（表格）", "@content"),
            ("課程照片（表格照片）", "@photos")
        ]
        
        # 新增報價單標記區塊
        self.price_list_tags = [
            ("品名（表格）", "@price_list_name"),
            ("單位（表格）", "@price_list_unit"),
            ("數量（表格）", "@price_list_quantity"),
            ("單價（表格）", "@price_list_price"),
            ("預計金額（表格）", "@price_list_amount"),
            ("用途說明（表格）", "@price_list_usage"),
            ("公司存摺（照片）", "@bank_account")
        ]
        
        # 初始化 teacher_combo
        self.teacher_combo = QComboBox()
        
        # 初始化資料管理器
        self.teacher_manager = TeacherDataManager()
        self.course_manager = CourseDataManager()
        
        # 建立資料夾
        self.create_directories()
        
        # 顯示登入視窗
        self.show_login_window()
    
    def create_directories(self):
        # 建立儲存圖片的資料夾
        os.makedirs('images/teachers', exist_ok=True)
        os.makedirs('images/courses', exist_ok=True)
    
    def show_login_window(self):
        self.login_widget = QWidget()
        self.setCentralWidget(self.login_widget)
        
        layout = QVBoxLayout()
        
        # 登入標籤
        label = QLabel("請輸入登入碼：")
        label.setFont(QFont('', 12))
        label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(label)
        
        # 登入輸入框
        self.login_input = QLineEdit()
        self.login_input.setEchoMode(QLineEdit.EchoMode.Password)
        self.login_input.setFixedWidth(200)
        layout.addWidget(self.login_input, alignment=Qt.AlignmentFlag.AlignCenter)
        
        # 登入按鈕
        self.login_button = QPushButton("登入")
        self.login_button.setFixedSize(200, 40)
        self.login_button.setFont(QFont('', 12))
        self.login_button.clicked.connect(self.verify_login)
        layout.addWidget(self.login_button, alignment=Qt.AlignmentFlag.AlignCenter)
        
        # 加入彈性空間
        layout.addStretch()
        
        # 加入版權聲明
        copyright_label = QLabel("Copyright © 2024 Zero Lin.")
        copyright_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        copyright_label.setStyleSheet("""
            QLabel {
                color: #666666;
                padding: 3px;
                font-size: 10px;
            }
        """)
        layout.addWidget(copyright_label)
        
        self.login_widget.setLayout(layout)
    
    def initialize_main_window(self):
        # 主視窗
        main_widget = QWidget()
        self.setCentralWidget(main_widget)
        
        # 主要布局
        main_layout = QVBoxLayout()
        
        # 分頁
        tab_widget = QTabWidget()
        
        # 標記工具分頁
        tags_tab = self.create_tags_tab()
        tab_widget.addTab(tags_tab, "標記工具")
        
        # 產生文件分頁
        doc_tab = self.create_doc_tab()
        tab_widget.addTab(doc_tab, "產生文件")
        
        main_layout.addWidget(tab_widget)
        
        # 加入版權聲明
        copyright_label = QLabel("Copyright © 2024 Zero Lin.")
        copyright_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        copyright_label.setStyleSheet("""
            QLabel {
                color: #666666;
                padding: 3px;
                font-size: 10px;
            }
        """)
        main_layout.addWidget(copyright_label)
        
        main_widget.setLayout(main_layout)
        
        # 更新教師列表
        self.update_teacher_list()
        
        # 設定視窗標和大小
        self.setWindowTitle("教師資料文件產生器")
        self.setFixedSize(400, 900)
        
        # 更新課程和教師列表
        self.update_course_list()
    
    def create_tags_tab(self):
        # 建立捲動區域
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        
        content_widget = QWidget()
        layout = QVBoxLayout()
        
        # 教師資料標記
        teacher_frame = QFrame()
        teacher_frame.setFrameStyle(QFrame.Shape.Box)
        teacher_layout = QVBoxLayout()
        
        teacher_label = QLabel("教師資料")
        teacher_label.setFont(QFont('', 12, QFont.Weight.Bold))
        teacher_layout.addWidget(teacher_label)
        
        # 使用網���布局來放置按鈕
        teacher_grid = QGridLayout()
        teacher_grid.setSpacing(10)  # 設定按鈕之間的間距
        teacher_grid.setHorizontalSpacing(10)  # 設定水平間距
        teacher_grid.setVerticalSpacing(10)    # 設定垂直間距
        
        # 計算按鈕寬度 (視窗寬度 - 邊距 - 中間間距) / 2
        button_width = (340 - 20 - 10) // 2
        
        for i, (label, tag) in enumerate(self.teacher_tags):
            btn = QPushButton(label)
            btn.setFixedHeight(40)
            btn.setFixedWidth(button_width)  # 設定固定寬度
            btn.clicked.connect(lambda checked, t=tag: self.copy_tag(t))
            row = i // 2
            col = i % 2
            teacher_grid.addWidget(btn, row, col)
        
        teacher_layout.addLayout(teacher_grid)
        teacher_frame.setLayout(teacher_layout)
        layout.addWidget(teacher_frame)
        
        # 課程資料標記區 (使用相同的設定)
        course_frame = QFrame()
        course_frame.setFrameStyle(QFrame.Shape.Box)
        course_layout = QVBoxLayout()
        
        course_label = QLabel("課程資料")
        course_label.setFont(QFont('', 12, QFont.Weight.Bold))
        course_layout.addWidget(course_label)
        
        course_grid = QGridLayout()
        course_grid.setSpacing(10)
        course_grid.setHorizontalSpacing(10)
        course_grid.setVerticalSpacing(10)
        
        for i, (label, tag) in enumerate(self.course_tags):
            btn = QPushButton(label)
            btn.setFixedHeight(40)
            btn.setFixedWidth(button_width)  # 使用相同寬度
            btn.clicked.connect(lambda checked, t=tag: self.copy_tag(t))
            row = i // 2
            col = i % 2
            course_grid.addWidget(btn, row, col)
        
        course_layout.addLayout(course_grid)
        course_frame.setLayout(course_layout)
        layout.addWidget(course_frame)
        
        # 報價單標記區 (使用相同的設定)
        price_list_frame = QFrame()
        price_list_frame.setFrameStyle(QFrame.Shape.Box)
        price_list_layout = QVBoxLayout()
        
        price_list_label = QLabel("報價單")
        price_list_label.setFont(QFont('', 12, QFont.Weight.Bold))
        price_list_layout.addWidget(price_list_label)
        
        price_list_grid = QGridLayout()
        price_list_grid.setSpacing(10)
        price_list_grid.setHorizontalSpacing(10)
        price_list_grid.setVerticalSpacing(10)
        
        for i, (label, tag) in enumerate(self.price_list_tags):
            btn = QPushButton(label)
            btn.setFixedHeight(40)
            btn.setFixedWidth(button_width)  # 使用相同的寬度
            btn.clicked.connect(lambda checked, t=tag: self.copy_tag(t))
            row = i // 2
            col = i % 2
            price_list_grid.addWidget(btn, row, col)
        
        price_list_layout.addLayout(price_list_grid)
        price_list_frame.setLayout(price_list_layout)
        layout.addWidget(price_list_frame)
        
        # 設定內容區域的邊距
        layout.setContentsMargins(10, 10, 10, 10)
        
        content_widget.setLayout(layout)
        scroll.setWidget(content_widget)
        return scroll
    
    def verify_login(self):
        print("開始驗證登入...")
        try:
            # 禁用登入按鈕並更改文字
            self.login_button.setEnabled(False)
            self.login_button.setText("資料更新中，請稍後...")
            QApplication.processEvents()  # 強制更新界面
            
            # 設定 scope
            self.SCOPES = [
                'https://www.googleapis.com/auth/spreadsheets',
                'https://www.googleapis.com/auth/drive.readonly'
            ]
            
            # 取得憑證
            print("取得服務帳號憑證...")
            self.creds = self.get_service_account_creds()
            
            # 建立 Sheets 服務
            print("建立 Sheets 服務...")
            self.service = build('sheets', 'v4', credentials=self.creds)
            
            # 建立登入管理器
            login_manager = LoginManager(self.creds)
            
            # 取得登入碼
            login_code = self.login_input.text()
            print(f"輸入的登入碼: {login_code}")
            
            # 驗證登入
            is_valid, need_update = login_manager.verify_login(login_code)
            
            if is_valid:
                print(f"登入成功，需要更新：{need_update}")
                if need_update:
                    print("開始更新資料...")
                    self.update_data()
                self.initialize_main_window()
            else:
                self.login_button.setEnabled(True)
                self.login_button.setText("登入")
                QMessageBox.critical(self, "錯誤", "登入碼無效")
                self.close()
            
        except Exception as e:
            self.login_button.setEnabled(True)
            self.login_button.setText("登入")
            error_msg = f"登入失敗：{str(e)}\n錯誤類型：{type(e)}"
            print(error_msg)
            print(f"錯誤詳情：{e.__dict__}")
            QMessageBox.critical(self, "錯誤", error_msg)
            self.close()
    
    def get_service_account_creds(self):
        try:
            import os
            import json
            
            # 先嘗試從環境變數讀取
            credentials_json = os.getenv('GOOGLE_CREDENTIALS')
            if credentials_json:
                try:
                    credentials_info = json.loads(credentials_json)
                    creds = service_account.Credentials.from_service_account_info(
                        credentials_info,
                        scopes=self.SCOPES
                    )
                    print("從環境變數成功載入憑證")
                    return creds
                except Exception as env_error:
                    print(f"從環境變數載入憑證失敗：{str(env_error)}")
            
            # 如果環境變數不存在或失敗，則從檔案讀取
            try:
                # 先嘗試讀取 embedded_credentials.json
                cred_file = 'embedded_credentials.json'
                if not os.path.exists(cred_file):
                    cred_file = 'credentials.json'
                
                print(f"嘗試從檔案讀取憑證：{cred_file}")
                with open(cred_file, 'r', encoding='utf-8') as f:
                    credentials_info = json.load(f)
                    creds = service_account.Credentials.from_service_account_info(
                        credentials_info,
                        scopes=self.SCOPES
                    )
                    print("從檔案成功載入憑證")
                    return creds
            except Exception as file_error:
                print(f"從檔案載入憑證失敗：{str(file_error)}")
                raise
        
        except Exception as e:
            print(f"讀取憑證時發生錯誤：{str(e)}")
            print(f"錯誤類型：{type(e)}")
            print(f"錯誤詳情：{e.__dict__}")
            raise
    
    def create_doc_tab(self):
        doc_widget = QWidget()
        layout = QVBoxLayout()
        
        # 課程選擇區域
        course_group = QFrame()
        course_group.setFrameStyle(QFrame.Shape.Box)
        course_layout = QVBoxLayout()
        
        course_label = QLabel("選擇課程")
        course_label.setFont(QFont('', 12, QFont.Weight.Bold))
        course_layout.addWidget(course_label)
        
        self.course_combo = QComboBox()
        self.course_combo.setFixedWidth(150)
        self.course_combo.currentIndexChanged.connect(self.update_teacher_list_by_course)
        course_layout.addWidget(self.course_combo)
        
        course_group.setLayout(course_layout)
        layout.addWidget(course_group)
        
        # 教師選擇區域
        teacher_group = QFrame()
        teacher_group.setFrameStyle(QFrame.Shape.Box)
        teacher_layout = QVBoxLayout()
        
        teacher_label = QLabel("選擇教師")
        teacher_label.setFont(QFont('', 12, QFont.Weight.Bold))
        teacher_layout.addWidget(teacher_label)
        
        self.teacher_combo = QComboBox()
        self.teacher_combo.setFixedWidth(150)
        teacher_layout.addWidget(self.teacher_combo)
        
        teacher_group.setLayout(teacher_layout)
        layout.addWidget(teacher_group)
        
        # 檔案選擇區域
        file_group = QFrame()
        file_group.setFrameStyle(QFrame.Shape.Box)
        file_layout = QVBoxLayout()
        
        file_label = QLabel("選擇範本")
        file_label.setFont(QFont('', 12, QFont.Weight.Bold))
        file_layout.addWidget(file_label)
        
        file_select_layout = QHBoxLayout()
        self.file_path = QLineEdit()
        self.file_path.setReadOnly(True)
        file_select_layout.addWidget(self.file_path)
        
        select_btn = QPushButton("瀏覽...")
        select_btn.clicked.connect(self.select_template)
        file_select_layout.addWidget(select_btn)
        
        file_layout.addLayout(file_select_layout)
        file_group.setLayout(file_layout)
        layout.addWidget(file_group)
        
        # 產生文件按鈕
        generate_btn = QPushButton("產生文件")
        generate_btn.setFixedHeight(50)
        generate_btn.setFont(QFont('', 12))
        generate_btn.clicked.connect(self.generate_document)
        layout.addWidget(generate_btn)
        
        layout.addStretch()
        doc_widget.setLayout(layout)
        return doc_widget
    
    def select_template(self):
        file_name, _ = QFileDialog.getOpenFileName(
            self,
            "選擇範本檔案",
            "",
            "Word Files (*.docx);;OpenDocument Files (*.odt)"
        )
        if file_name:
            self.file_path.setText(file_name)
    
    def generate_document(self):
        try:
            print("開始生文件...")
            
            # 檢查選擇
            if not self.teacher_combo.currentText() or not self.course_combo.currentText():
                print("未選擇課程或教師")
                QMessageBox.warning(self, "警告", "請選擇課程和教師")
                return
            
            if not self.file_path.text():
                print("未選擇範本檔案")
                QMessageBox.warning(self, "警告", "請選擇範本檔案")
                return
            
            print("取得教師和課程ID...")
            teacher_id = self.teacher_combo.currentData()
            course_id = self.course_combo.currentData()
            
            print(f"教師ID: {teacher_id}")
            print(f"課程ID: {course_id}")
            
            # 取得教師和課程資料
            print("讀取教師資料...")
            teacher_data = self.teacher_manager.teachers.get(teacher_id)
            if not teacher_data:
                print(f"找不到教師資料: {teacher_id}")
                raise ValueError(f"找不到教師資料: {teacher_id}")
            
            print("讀取課程資料...")
            course_data = self.course_manager.courses.get(course_id)
            if not course_data:
                print(f"找不到課程資料: {course_id}")
                raise ValueError(f"找不到課程資料: {course_id}")
            
            print("併資料...")
            combined_data = {**teacher_data, **course_data}
            
            # 使用範本檔的目錄
            template_dir = os.path.dirname(self.file_path.text())
            print(f"範本目錄: {template_dir}")
            
            # 建立新的檔名
            output_filename = f"{course_data.get('course_name', 'unknown')} - {teacher_data.get('name', 'unknown')}.docx"
            output_path = os.path.join(template_dir, output_filename)
            print(f"輸出檔案路徑: {output_path}")
            
            # 處理文件
            print("開始處理文件...")
            processor = DocumentProcessor(self.file_path.text())
            output_path = processor.process_document(combined_data, output_path)
            
            print("文件成完成")
            QMessageBox.information(self, "成功", f"文件已產生：{output_path}")
            
        except Exception as e:
            error_msg = f"產生文件失敗：{str(e)}\n錯誤類型：{type(e)}"
            print(error_msg)
            print(f"錯誤詳情：{e.__dict__}")
            QMessageBox.critical(self, "錯誤", error_msg)
    
    def _download_and_save_photo(self, drive_service, folder_id, photo_name, save_path):
        try:
            print(f"開始處理照片：{photo_name}")
            
            # 搜尋特定檔案（考慮常見的圖片副檔）
            possible_extensions = ['.jpg', '.jpeg', '.png', '.gif']
            
            # 如果檔名已經包含副檔名
            if any(photo_name.lower().endswith(ext) for ext in possible_extensions):
                query = f"name = '{photo_name}' and '{folder_id}' in parents and trashed = false"
            else:
                # 如果檔名不包含副檔名，使用 OR 條件搜尋所有可能的副檔名
                extension_conditions = [f"name = '{photo_name}{ext}'" for ext in possible_extensions]
                query = f"({' or '.join(extension_conditions)}) and '{folder_id}' in parents and trashed = false"
            
            results = drive_service.files().list(
                q=query,
                fields="files(id, name)",
                pageSize=1
            ).execute()
            
            matching_files = results.get('files', [])
            
            if not matching_files:
                print(f"找不到檔案：{photo_name}")
                return
            
            file_id = matching_files[0]['id']
            file_name = matching_files[0]['name']
            
            # 更新儲存路徑以含正確的副檔名
            save_path = os.path.splitext(save_path)[0] + os.path.splitext(file_name)[1]
            
            # 下載檔案
            request = drive_service.files().get_media(fileId=file_id)
            fh = io.BytesIO()
            downloader = MediaIoBaseDownload(fh, request)
            
            done = False
            while done is False:
                status, done = downloader.next_chunk()
            
            # 儲存檔案
            fh.seek(0)
            os.makedirs(os.path.dirname(save_path), exist_ok=True)
            with open(save_path, 'wb') as f:
                f.write(fh.read())
            
        except Exception as e:
            print(f"下載照片失敗 {photo_name}: {str(e)}")
            if hasattr(e, 'resp') and hasattr(e.resp, 'status'):
                print(f"HTTP狀態碼：{e.resp.status}")
    
    def update_teacher_list(self):
        self.teacher_combo.clear()
        for tid, tdata in self.teacher_manager.teachers.items():
            self.teacher_combo.addItem(f"{tid}: {tdata['name']}")
    
    def copy_tag(self, tag):
        # 複製標記剪貼簿
        QApplication.clipboard().setText(tag)
        
        # 立浮動提示視窗
        toast = QLabel(f"已複製 {tag}", self)
        toast.setFont(QFont('', 12))
        toast.setStyleSheet("""
            QLabel {
                background-color: #4CAF50;
                color: white;
                padding: 10px 20px;
                border-radius: 5px;
            }
        """)
        toast.setAlignment(Qt.AlignmentFlag.AlignCenter)
        
        # 調整提示視窗大小
        toast.adjustSize()
        
        # 計算位置（置中顯示）
        pos_x = (self.width() - toast.width()) // 2
        pos_y = self.height() - toast.height() - 50  # 離底部 50 像素
        toast.move(pos_x, pos_y)
        
        # 顯示提示
        toast.show()
        
        # 1.5秒後自動移除提示
        QTimer.singleShot(1500, toast.deleteLater)
    
    def toggle_topmost(self):
        current = self.isActiveWindow()
        self.setWindowFlags(self.windowFlags() | Qt.WindowStaysOnTopHint if not current else self.windowFlags() & ~Qt.WindowStaysOnTopHint)
    
    def update_data(self):
        try:
            print("開始更新資料...")
            
            # 清除所有現有資料
            self.clear_all_data()
                       
            # 更新師資資料
            self.import_teacher_data_from_google()
            print("師資資料更新成")
            
            # 更新課程資料
            self.import_course_data_from_google()
            print("課程資料更新完成")
            
        except Exception as e:
            print(f"更新過程發生錯誤：{str(e)}")
            QMessageBox.critical(self, "錯誤", f"更新失敗：{str(e)}")
    
    def clear_all_data(self):
        # 清除 JSON 檔案
        if os.path.exists('teachers.json'):
            os.remove('teachers.json')
        if os.path.exists('courses.json'):
            os.remove('courses.json')
        
        # 清除圖片資料夾
        if os.path.exists('images'):
            shutil.rmtree('images')
        os.makedirs('images/teachers', exist_ok=True)
        os.makedirs('images/courses', exist_ok=True)
        
        # 重置資料管理器
        self.teacher_manager = TeacherDataManager()
        self.course_manager = CourseDataManager()
    
    def import_teacher_data_from_google(self):
        try:
            print("取師資料...")
            # 取得使用者地區
            login_sheet = self.service.spreadsheets().values().get(
                spreadsheetId=self.SPREADSHEET_ID,
                range='login!B2:B'
            ).execute()
            user_region = login_sheet.get('values', [[None]])[0][0]
            print(f"使用者地區：{user_region}")
            
            # 讀取資資料
            result = self.service.spreadsheets().values().get(
                spreadsheetId=self.SPREADSHEET_ID,
                range='師資!A2:V'
            ).execute()
            values = result.get('values', [])
            
            # 清除現有資料
            if os.path.exists('teachers.json'):
                os.remove('teachers.json')
            if os.path.exists('images/teachers'):
                shutil.rmtree('images/teachers')
                os.makedirs('images/teachers')
            
            # 建立 Drive 服務
            drive_service = build('drive', 'v3', credentials=self.creds)
            # 確保資料夾存在
            if not os.path.exists('images/teachers'):
                os.makedirs('images/teachers')
            if not os.path.exists('images/courses'):
                os.makedirs('images/courses')

            # 設正確的資料夾 ID（確認這是正確的資料夾 ID）
            folder_id = "1I_LHNBh8pDMJRbtRmQRblqumIZ9vayfa"
            
            # 處理每一筆教師資料
            for row in values:
                try:
                    if len(row) > 0 and row[0] != user_region:
                        continue
                    
                    # 確保姓名欄位存在
                    name = row[1] if len(row) > 1 else ''
                    
                    # 使用預設值處理所有欄位
                    teacher_data = {
                        'region': row[0] if len(row) > 0 else '',      # 地區 (A欄)
                        'name': row[1] if len(row) > 1 else '',        # 姓名 (B欄)
                        'nickname': row[2] if len(row) > 2 else '',    # 綽號 (C欄)
                        'photo': '',                                   # 大頭照 (D欄)
                        'unit': row[4] if len(row) > 4 else '',       # 申請單位 (E欄)
                        'birth': row[5] if len(row) > 5 else '',      # 出生日期 (F欄)
                        'gender': row[6] if len(row) > 6 else '',     # 性別 (G欄)
                        'mobile': row[7] if len(row) > 7 else '',     # 手機 (H欄)
                        'idno': row[8] if len(row) > 8 else '',       # 身分證字號 (I欄)
                        'address': row[9] if len(row) > 9 else '',    # 通訊地址 (J欄)
                        'email': row[10] if len(row) > 10 else '',    # Email (K欄)
                        'line': row[11] if len(row) > 11 else '',     # Line ID (L欄)
                        'skill': row[12] if len(row) > 12 else '',    # 專長 (M欄)
                        'experience': row[13] if len(row) > 13 else '',# 教學經驗 (N欄)
                        'history': row[14] if len(row) > 14 else '',  # 經歷 (O欄)
                        'education': row[15] if len(row) > 15 else '', # 最高學歷 (P欄)
                        'job': row[16] if len(row) > 16 else '',      # 現職 (Q欄)
                        'id_front': '',                               # 身分證正面 (R欄)
                        'id_back': '',                                # 身分證反面 (S欄)
                        'diploma': '',                                # 畢業證書 (T欄)
                        'other_certs': [],                           # 其他證明 (U欄)
                        'course_type': row[21] if len(row) > 21 else '' # 課程分類 (V欄)
                    }
                    
                    # 修改照片欄位對應
                    photo_fields = {
                        3: ('photo', '大頭照'),       # D欄
                        17: ('id_front', '身分證正'),   # S欄
                        18: ('id_back', '身分證反'),    # T欄
                        19: ('diploma', '畢業證書')     # U欄
                    }
                    
                    # 修改這部分，加入錯誤處理
                    for col_idx, (field_name, photo_type) in photo_fields.items():
                        try:
                            if len(row) > col_idx and row[col_idx]:
                                try:
                                    # 使用教師姓名和照片類型建立檔名
                                    photo_name = f"{row[1]}{photo_type}"  # 例如：林顯庭大頭照
                                    save_path = f'images/teachers/{photo_name}'
                                    self._download_and_save_photo(
                                        drive_service, 
                                        folder_id, 
                                        photo_name, 
                                        save_path
                                    )
                                    teacher_data[field_name] = save_path
                                except Exception as photo_error:
                                    print(f"跳過照片 {field_name}: {str(photo_error)}")
                                    teacher_data[field_name] = ''
                        except Exception as field_error:
                            print(f"處理照片欄位 {field_name} 時發生錯誤: {str(field_error)}")
                            teacher_data[field_name] = ''
                    
                    # 處理其他證明
                    try:
                        teacher_name = row[1]
                        # 修改搜尋條件，使用 contains 和 and 組合
                        query = f"name contains '{teacher_name}其他證明_' and '{folder_id}' in parents and trashed = false"
                        results = drive_service.files().list(
                            q=query,
                            fields="files(id, name)",
                            orderBy="name"  # 加入排序，確保檔案順序
                        ).execute()
                        
                        other_certs = []
                        for file in results.get('files', []):
                            try:
                                photo_name = file['name']
                                save_path = f'images/teachers/{photo_name}'
                                self._download_and_save_photo(drive_service, folder_id, photo_name, save_path)
                                other_certs.append(save_path)
                            except Exception as cert_error:
                                print(f"跳過其他證明照片 {photo_name}: {str(cert_error)}")
                                continue
                        
                        teacher_data['other_certs'] = other_certs
                    except Exception as other_certs_error:
                        print(f"處理其他證明時發生錯誤: {str(other_certs_error)}")
                        teacher_data['other_certs'] = []
                    
                    # 只有當姓名不為空時才新增教師資料
                    if name:
                        self.teacher_manager.add_teacher(teacher_data)
                    else:
                        print(f"跳過沒有姓名的資料")
                    
                except Exception as row_error:
                    print(f"處理教師資料時發生錯誤，跳過此筆資料: {str(row_error)}")
                    continue
            
            self.update_teacher_list()
            
        except Exception as e:
            print(f"更新師資資料時發生錯誤：{str(e)}")
            QMessageBox.critical(self, "錯誤", f"更新失敗：{str(e)}")

    def import_course_data_from_google(self):
        try:
            print("讀取課程資料...")
            result = self.service.spreadsheets().values().get(
                spreadsheetId=self.SPREADSHEET_ID,
                range='課程!A2:P'  # 修改範圍到 O 欄
            ).execute()
            values = result.get('values', [])
            print(f"讀取到 {len(values)} 筆課程資料")
            
            # 清除現有資料
            if os.path.exists('courses.json'):
                os.remove('courses.json')
            if os.path.exists('images/courses'):
                shutil.rmtree('images/courses')
                os.makedirs('images/courses')
            
            # 建立 Drive 服務
            drive_service = build('drive', 'v3', credentials=self.creds)
            # 確保資料夾存在
            if not os.path.exists('images/teachers'):
                os.makedirs('images/teachers')
            if not os.path.exists('images/courses'):
                os.makedirs('images/courses')

            # 設定正確的資料夾 ID（請確認這是正確的料夾 ID）
            folder_id = "1I_LHNBh8pDMJRbtRmQRblqumIZ9vayfa"
            
            # 處理每一筆課程資料
            for row in values:
                try:
                    course_data = {
                        'course_name': row[0] if len(row) > 0 else '',     # 社團名稱 (A欄)
                        'intro': row[1] if len(row) > 1 else '',           # 課程介紹 (B欄)
                        'material_fee': row[2] if len(row) > 2 else '',    # 材料費 (C欄)
                        'reason': row[3] if len(row) > 3 else '',          # 原因 (D欄)
                        'target': row[4] if len(row) > 4 else '',          # 教學目標 (E欄)
                        'course_topic': row[5] if len(row) > 5 else '',    # 課程主題 (F欄) - 新增
                        'content': row[6] if len(row) > 6 else '',         # 課程內容 (G欄)
                        'photos': [],                                      # 課程照片 (H欄)
                        'price_list_name': row[8] if len(row) > 8 else '', # 報價單品名 (I欄)
                        'price_list_unit': row[9] if len(row) > 9 else '', # 報價單單位 (J欄)
                        'price_list_quantity': row[10] if len(row) > 10 else '', # 報價單數量 (K欄)
                        'price_list_price': row[11] if len(row) > 11 else '',    # 報價單單價 (L欄)
                        'price_list_amount': row[12] if len(row) > 12 else '',   # 報價單預計金額 (M欄)
                        'price_list_usage': row[13] if len(row) > 13 else '',    # 報價單用途說明 (N欄)
                        'bank_account': '',                                      # 公司存摺 (O欄)
                        'course': row[15] if len(row) > 15 else ''              # 課程分類 (P欄)
                    }
                    
                    # 處理課程照片
                    if len(row) > 7:
                        try:
                            # 從課程名稱建立搜尋條件
                            course_name = row[15].strip()  # 使用 A 欄的課程名稱
                            # 修改搜尋條件，使用 contains 和 and 組合
                            query = f"name contains '{course_name}_' and '{folder_id}' in parents and trashed = false"
                            results = drive_service.files().list(
                                q=query,
                                fields="files(id, name)",
                                orderBy="name"  # 加入排序，確保檔案順序
                            ).execute()
                            
                            for file in results.get('files', []):
                                try:
                                    photo_name = file['name']
                                    # 確認檔名格式是否符合 課程名稱_數字 的格式
                                    name_parts = photo_name.rsplit('_', 1)
                                    if len(name_parts) == 2 and name_parts[1].split('.')[0].isdigit():
                                        save_path = f'images/courses/{photo_name}'
                                        self._download_and_save_photo(drive_service, folder_id, photo_name, save_path)
                                        course_data['photos'].append(save_path)
                                except Exception as photo_error:
                                    print(f"跳過照片 {photo_name}: {str(photo_error)}")
                        except Exception as split_error:
                            print(f"處理照片列表時發生錯誤: {str(split_error)}")
                    
                    # 處司存摺（O欄）
                    if len(row) > 14 and row[14]:
                        try:
                            save_path = f'images/courses/{row[14]}'
                            self._download_and_save_photo(drive_service, folder_id, row[14], save_path)
                            course_data['bank_account'] = save_path
                        except Exception as bank_error:
                            print(f"跳過公司存摺照片: {str(bank_error)}")
                            course_data['bank_account'] = ''
                    
                    self.course_manager.add_course(course_data)
                    
                except Exception as row_error:
                    print(f"處理課程資料時發生錯誤，跳過此筆資料: {str(row_error)}")
                    continue
                        
        except Exception as e:
            print(f"更新課程資料時發生錯誤：{str(e)}")
            QMessageBox.critical(self, "錯誤", f"更新失敗：{str(e)}")

    def update_course_list(self):
        self.course_combo.clear()
        
        # 收集所有不重複的課程分類
        course_types = {}  # 改用字典來儲存課分類和ID的對應關係
        for course_id, cdata in self.course_manager.courses.items():
            if 'course' in cdata and cdata['course']:
                course_types[cdata['course']] = course_id
        
        # 將課程分類加入下拉選單
        for course_type, course_id in sorted(course_types.items()):
            self.course_combo.addItem(course_type, course_id)  # 使用課程ID作為數據

    def update_teacher_list_by_course(self):
        try:
            selected_course = self.course_combo.currentText()
            self.teacher_combo.clear()
            
            if not selected_course:
                return
            
            # 取得所有教師資料
            with open('teachers.json', 'r', encoding='utf-8') as f:
                teachers = json.load(f)
            
            # 篩選符合課程的教師
            matching_teachers = []
            for teacher_id, teacher in teachers.items():
                # 將課程類型拆分為列表
                course_types = [ct.strip() for ct in teacher.get('course_type', '').split('、')]
                # 如果選擇的課程在教師的課程類型列表中
                if selected_course in course_types:
                    matching_teachers.append((teacher_id, teacher['name']))
            
            # 將符合的教師加入下拉單
            for teacher_id, name in matching_teachers:
                self.teacher_combo.addItem(name, userData=teacher_id)
            
        except Exception as e:
            print(f"更新教師列表時發生錯誤：{str(e)}")
            QMessageBox.critical(self, "錯誤", f"更新教師列表時發生錯誤：{str(e)}")

    def verify_drive_api(self):
        try:
            drive_service = build('drive', 'v3', credentials=self.creds)
            # 測試 API 是否正常運作
            drive_service.files().list(pageSize=1).execute()
            print("Google Drive API 驗證成功")
            return True
        except Exception as e:
            print(f"Google Drive API 驗證失敗：{str(e)}")
            return False

class TeacherDataManager:
    def __init__(self):
        self.teachers = self._load_teachers()
    
    def _load_teachers(self):
        if os.path.exists('teachers.json'):
            with open('teachers.json', 'r', encoding='utf-8') as f:
                return json.load(f)
        return {}
    
    def add_teacher(self, teacher_data):
        teacher_id = str(len(self.teachers) + 1)
        self.teachers[teacher_id] = teacher_data
        self._save_teachers()
        
    def _save_teachers(self):
        with open('teachers.json', 'w', encoding='utf-8') as f:
            json.dump(self.teachers, f, ensure_ascii=False, indent=2)

class CourseDataManager:
    def __init__(self):
        self.courses = self._load_courses()
    
    def _load_courses(self):
        if os.path.exists('courses.json'):
            with open('courses.json', 'r', encoding='utf-8') as f:
                return json.load(f)
        return {}
    
    def add_course(self, course_data):
        course_id = str(len(self.courses) + 1)
        self.courses[course_id] = course_data
        self._save_courses()
    
    def _save_courses(self):
        with open('courses.json', 'w', encoding='utf-8') as f:
            json.dump(self.courses, f, ensure_ascii=False, indent=2)

class DocumentProcessor:
    def __init__(self, template_path):
        self.template_path = template_path
        self.file_type = os.path.splitext(template_path)[1].lower()
    
    def process_document(self, data, output_path):
        try:
            print("開始處理文件...")
            doc = Document(self.template_path)
            
            # 修改這裡：在檔名中加入時間戳記，確保檔名唯一
            timestamp = datetime.now().strftime('%Y%m%d')
            output_filename = f"{data.get('course_name', 'unknown')} - {data.get('name', 'unknown')}_{timestamp}.docx"
            output_path = os.path.join(os.path.dirname(self.template_path), output_filename)
            
            # 處理表格中的標記
            print("處理表格內容...")
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if '@content' in cell.text:
                            self._replace_course_table(table, data)
                            break
                        elif '@course_topic' in cell.text:  # 新增這個條件
                            self._replace_course_topic_table(table, data)
                            break
                        elif '@price_list_table' in cell.text:
                            self._replace_price_list_table(table, data)
                            break
                        elif '@photos' in cell.text:
                            self._replace_photos_table(table, data)
                            break
                        else:
                            self._process_cell(cell, data)
            
            # 處理段落中的標記
            print("處理段標記內容...")
            for paragraph in doc.paragraphs:
                self._process_paragraph(paragraph, data)
            
            print(f"儲存文件到 {output_path}...")
            doc.save(output_path)
            return output_path
        
        except Exception as e:
            print(f"處理文件時發生錯誤：{str(e)}")
            raise
    
    def _process_cell(self, cell, data):
        try:
            for key, value in data.items():
                marker = f"@{key}"
                if marker in cell.text:
                    print(f"處理標記 {marker}...")
                    if isinstance(value, list) and key == 'other_certs':
                        # 處理多張圖片
                        print(f"開始處理其他證明照片，共 {len(value)} 張")
                        cell.text = cell.text.replace(marker, '')
                        paragraph = cell.paragraphs[0]
                        
                        for img_path in value:
                            print(f"處理照片：{img_path}")
                            # 檢查檔案是否存在（包含各種可能的副檔名）
                            possible_extensions = ['.jpg', '.jpeg', '.png', '.gif']
                            full_path = None
                            
                            # 先檢查原始路徑
                            if os.path.exists(img_path):
                                full_path = img_path
                            else:
                                # 嘗試不同的副檔名
                                base_path = os.path.splitext(img_path)[0]
                                for ext in possible_extensions:
                                    test_path = f"{base_path}{ext}"
                                    if os.path.exists(test_path):
                                        full_path = test_path
                                        break
                            
                            if full_path and os.path.exists(full_path):
                                print(f"插入照片：{full_path}")
                                run = paragraph.add_run()
                                run.add_picture(full_path, width=Inches(2))
                                run.add_text('\n')  # 在每張照片後添加換行
                            else:
                                print(f"找不到照片：{img_path}")
                                
                    elif key in ['photo', 'id_front', 'id_back', 'diploma', 'bank_account'] and value:
                        # 處理單張圖片
                        cell.text = cell.text.replace(marker, '')
                        paragraph = cell.paragraphs[0]
                        
                        # 檢查檔案是否存在（包含各種可能的副檔名）
                        possible_extensions = ['.jpg', '.jpeg', '.png', '.gif']
                        full_path = None
                        for ext in possible_extensions:
                            test_path = f"{value}{ext}"
                            if os.path.exists(test_path):
                                full_path = test_path
                                break
                        
                        if full_path and os.path.exists(full_path):
                            run = paragraph.add_run()
                            run.add_picture(full_path, width=Inches(2))
                    else:
                        # 處理文字
                        cell.text = cell.text.replace(marker, str(value) if value else '')
        except Exception as e:
            print(f"處理儲存格時發生錯誤：{str(e)}")
            print(f"錯誤類型：{type(e)}")
            print(f"錯誤詳情：{e.__dict__}")
            raise
    
    def _process_paragraph(self, paragraph, data):
        try:
            for key, value in data.items():
                marker = f"@{key}"
                if marker in paragraph.text:
                    print(f"處理標記 {marker}...")
                    if isinstance(value, list) and key == 'other_certs':
                        # 處理多張圖片
                        paragraph.text = paragraph.text.replace(marker, '')
                        for img_path in value:
                            # 檢查檔案是否存在（包含各種可能的副檔名）
                            possible_extensions = ['.jpg', '.jpeg', '.png', '.gif']
                            full_path = None
                            for ext in possible_extensions:
                                test_path = f"{img_path}{ext}"
                                if os.path.exists(test_path):
                                    full_path = test_path
                                    break
                            
                            if full_path and os.path.exists(full_path):
                                run = paragraph.add_run()
                                run.add_picture(full_path, width=Inches(2))
                                run.add_text('\n')
                    elif key in ['photo', 'id_front', 'id_back', 'diploma', 'bank_account'] and value:
                        # 處理單張圖片
                        paragraph.text = paragraph.text.replace(marker, '')
                        
                        # 檢查檔案是否存在（包含各種可能的副檔名）
                        possible_extensions = ['.jpg', '.jpeg', '.png', '.gif']
                        full_path = None
                        for ext in possible_extensions:
                            test_path = f"{value}{ext}"
                            if os.path.exists(test_path):
                                full_path = test_path
                                break
                        
                        if full_path and os.path.exists(full_path):
                            run = paragraph.add_run()
                            run.add_picture(full_path, width=Inches(2))
                    else:
                        # 處理文字
                        paragraph.text = paragraph.text.replace(marker, str(value) if value else '')
        except Exception as e:
            print(f"處理段落時發生錯誤：{str(e)}")
            raise
    
    def _replace_course_table(self, table, data):
        try:
            print("開始處理課程內容表格...")
            
            # 找到標記所在的儲存格位
            start_row = 0
            start_col = 0
            found = False
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    if '@content' in cell.text:
                        start_row = i
                        start_col = j
                        found = True
                        # 清除標記
                        cell.text = cell.text.replace('@content', '')
                        break
                if found:
                    break
            
            if not found:
                print("找不到課程內容表格標記")
                return
            
            # 取得課程內容
            content = data.get('content', '')
            if not content:
                print("找不到課程內容")
                return
            
            # 將課程內容分行
            content_lines = content.split('\n')
            print(f"課程內共 {len(content_lines)} 行")
            
            # 確保表格有足夠的列數
            current_row = start_row
            for content_line in content_lines:
                if current_row >= len(table.rows):
                    table.add_row()
                
                # 從指定的欄位開始填入
                if start_col < len(table.rows[current_row].cells):
                    table.rows[current_row].cells[start_col].text = content_line.strip()
                current_row += 1
            
            print("課程內容表格處理完成")
            
        except Exception as e:
            print(f"處理課程內容表格時發生錯誤：{str(e)}")
            raise
    
    def _replace_price_list_table(self, table, data):
        try:
            print("開始處理報價單表格...")
            
            # 找到標記所在的儲存格位
            start_row = 0
            start_col = 0
            found = False
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    if '@price_list_table' in cell.text:
                        start_row = i
                        start_col = j
                        found = True
                        # 清除標記
                        cell.text = cell.text.replace('@price_list_table', '')
                        break
                if found:
                    break
            
            if not found:
                print("找不到報價單表格標記")
                return
            
            # 報價單欄位和對應的資料
            price_list_items = [
                ('price_list_name', '品項'),
                ('price_list_unit', '單位'),
                ('price_list_quantity', '數量'),
                ('price_list_price', '單價'),
                ('price_list_amount', '預計金額'),
                ('price_list_usage', '用途說明')
            ]
            
            # 填入報價單資料
            current_row = start_row
            for field, label in price_list_items:
                if current_row < len(table.rows):
                    value = data.get(field, '')
                    # 從指定的欄位開始填
                    if start_col < len(table.rows[current_row].cells):
                        table.rows[current_row].cells[start_col].text = str(value)
                    current_row += 1  # 移到下一行
            
            print("報價單表格處理完成")
            
        except Exception as e:
            print(f"處理報價單表格時發生錯誤：{str(e)}")
            raise
    
    def _replace_photos_table(self, table, data):
        try:
            print("開始處理課程照片表格...")
            
            # 找到標記所在的儲存格位置
            start_row = 0
            start_col = 0
            found = False
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    if '@photos' in cell.text:
                        start_row = i
                        start_col = j
                        found = True
                        # 清除標記
                        cell.text = cell.text.replace('@photos', '')
                        break
                if found:
                    break
            
            if not found:
                print("找不到課程照片表格標記")
                return
            
            # 取得照片列表
            photos = data.get('photos', [])
            if not photos:
                print("找不到課程照片")
                return
            
            # 解析照片資訊
            photo_info = {}  # 用來儲存週次和對應的照片路徑
            for photo_path in photos:
                # 從檔名中提取週次資訊
                base_name = os.path.basename(photo_path)
                name_without_ext = os.path.splitext(base_name)[0]
                
                # 假設檔名格式為 "課程名稱_N"
                try:
                    week_num = int(name_without_ext.split('_')[-1])
                    photo_info[week_num] = photo_path
                except (ValueError, IndexError):
                    print(f"無法從檔名 {base_name} 解析週次資訊")
                    continue
            
            # 填入照片
            current_row = start_row
            week = 1
            while current_row < len(table.rows):
                if week in photo_info:
                    photo_path = photo_info[week]
                    # 檢查檔案是否存在（包含各種可能的副檔名）
                    possible_extensions = ['.jpg', '.jpeg', '.png', '.gif']
                    full_path = None
                    
                    # 先檢查原始路徑
                    if os.path.exists(photo_path):
                        full_path = photo_path
                    else:
                        # 嘗試不同的副檔名
                        base_path = os.path.splitext(photo_path)[0]
                        for ext in possible_extensions:
                            test_path = f"{base_path}{ext}"
                            if os.path.exists(test_path):
                                full_path = test_path
                                break
                    
                    if full_path and os.path.exists(full_path):
                        cell = table.rows[current_row].cells[start_col]
                        cell.text = ''  # 清空儲存格
                        paragraph = cell.paragraphs[0]
                        run = paragraph.add_run()
                        run.add_picture(full_path, width=Inches(2))
                
                current_row += 1
                week += 1
            
            print("課程照片表格處理完成")
            
        except Exception as e:
            print(f"處理課程照片表格時發生錯誤：{str(e)}")
            raise
    
    def _replace_course_topic_table(self, table, data):
        try:
            print("開始處理課程主題表格...")
            
            # 找到標記所在的儲存格位置
            start_row = 0
            start_col = 0
            found = False
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    if '@course_topic' in cell.text:
                        start_row = i
                        start_col = j
                        found = True
                        # 清除標記
                        cell.text = cell.text.replace('@course_topic', '')
                        break
                if found:
                    break
            
            if not found:
                print("找不到課程主題表格標記")
                return
            
            # 取得課程主題內容
            topic = data.get('course_topic', '')
            if not topic:
                print("找不到課程主題內容")
                return
            
            # 將課程主題分行
            topic_lines = topic.split('\n')
            print(f"課程主題共 {len(topic_lines)} 行")
            
            # 確保表格有足夠的列數
            current_row = start_row
            for topic_line in topic_lines:
                if current_row >= len(table.rows):
                    table.add_row()
                
                # 從指定的欄位開始填入
                if start_col < len(table.rows[current_row].cells):
                    table.rows[current_row].cells[start_col].text = topic_line.strip()
                current_row += 1
            
            print("課程主題表格處理完成")
            
        except Exception as e:
            print(f"處理課程主題表格時發生錯誤：{str(e)}")
            raise

def main():
    import sys
    
    try:
        app = QApplication(sys.argv)
        
        # 設定應用程式樣式
        app.setStyle('Fusion')
        
        # 建立並顯示主視窗
        window = TeacherDocApp()
        window.show()
        
        # 執行應用程式
        sys.exit(app.exec())
        
    except Exception as e:
        print(f"程式執行錯誤：{str(e)}")
        sys.exit(1)

if __name__ == "__main__":
    main()
